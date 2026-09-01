"""
Script: create_lkm_lengkap.py
Membuat Lembar Kegiatan Mahasiswa (LKM) yang lengkap dengan gambar, tabel,
dan konten detail berdasarkan penelitian Deep Learning Rupiah Prediction.
"""

import os
import sys

# Auto-install python-docx jika belum ada
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx
except ImportError:
    print("Menginstal python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_bg(cell, hex_color):
    """Mengatur warna latar belakang sel tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Mengatur border sel tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        for key, value in kwargs.items():
            element.set(qn('w:{}'.format(key)), value)
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None, center=False):
    """Menambahkan heading dengan styling kustom."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(14)
        if color is None:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Navy Blue
    elif level == 2:
        run.font.size = Pt(12)
        if color is None:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)  # Blue
    elif level == 3:
        run.font.size = Pt(11)
        if color is None:
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    if color:
        run.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(doc, text, indent=False, bold=False, italic=False, size=10.5):
    """Menambahkan paragraf body text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_bullet(doc, text, level=1, bold_prefix=None):
    """Menambahkan bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.bold = True
        run_bold.font.size = Pt(10.5)
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.3 * level)
    return p

def add_numbered(doc, text, bold_prefix=None, size=10.5):
    """Menambahkan nomor list."""
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.bold = True
        run_bold.font.size = Pt(size)
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(size)
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
    return p

def add_figure(doc, img_path, caption, width=6.0):
    """Menambahkan gambar dengan caption di bawahnya."""
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(caption)
        run_cap.font.italic = True
        run_cap.font.size = Pt(9.5)
        run_cap.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    else:
        p = doc.add_paragraph(f"[Gambar tidak ditemukan: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_separator(doc):
    """Menambahkan garis horizontal."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_info_box(doc, label, value):
    """Menambahkan info box dalam format tabel 2 kolom."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    cell_label = table.cell(0, 0)
    cell_value = table.cell(0, 1)
    cell_label.width = Inches(2.0)
    cell_value.width = Inches(4.5)
    set_cell_bg(cell_label, '2E75B6')
    set_cell_bg(cell_value, 'EBF3FB')
    run_label = cell_label.paragraphs[0].add_run(label)
    run_label.font.bold = True
    run_label.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run_label.font.size = Pt(10)
    cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_value = cell_value.paragraphs[0].add_run(value)
    run_value.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ============================================================
# MAIN DOCUMENT CREATION
# ============================================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(BASE_DIR, "LKM_Bisnis_Inteligent_LENGKAP.docx")

print("Membuat dokumen LKM yang lengkap...")
doc = Document()

# ---- PAGE MARGINS ----
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

section = doc.sections[0]
section.page_width = Inches(8.27)   # A4
section.page_height = Inches(11.69) # A4
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ============================================================
# HEADER - JUDUL UTAMA
# ============================================================
p_lkm = doc.add_paragraph()
p_lkm.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_lkm = p_lkm.add_run("LEMBAR KEGIATAN MAHASISWA (LKM)")
r_lkm.font.bold = True
r_lkm.font.size = Pt(16)
r_lkm.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p_mk = doc.add_paragraph()
p_mk.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_mk = p_mk.add_run("Mata Kuliah: Bisnis Inteligent")
r_mk.font.bold = True
r_mk.font.size = Pt(12)

add_separator(doc)
doc.add_paragraph()

# ============================================================
# IDENTITAS LKM - Tabel Info
# ============================================================
add_heading(doc, "IDENTITAS KEGIATAN", level=1)
doc.add_paragraph()

id_table = doc.add_table(rows=9, cols=2)
id_table.style = 'Table Grid'
id_table.alignment = WD_TABLE_ALIGNMENT.CENTER

id_data = [
    ("Nama Mata Kuliah", "Bisnis Inteligent"),
    ("Kode Mata Kuliah", "MBI-502"),
    ("Program Studi", "Magister Manajemen (S2)"),
    ("Institusi", "Universitas Nusa Putra"),
    ("Semester", "Semester 2 (Genap 2025/2026)"),
    ("Topik Penelitian", "Analisis dan Prediksi Nilai Tukar Rupiah (USD/IDR) Menggunakan Deep Learning Berbasis Indikator Makroekonomi"),
    ("Waktu Kegiatan", "±6 Jam (Teori + Praktikum)"),
    ("Jenis Kegiatan", "Praktikum Individual / Kelompok"),
]

header_row = id_table.rows[0]
set_cell_bg(header_row.cells[0], '1F497D')
set_cell_bg(header_row.cells[1], '1F497D')
for i, cell in enumerate(header_row.cells):
    headers = ["KOMPONEN", "KETERANGAN"]
    run_h = cell.paragraphs[0].add_run(headers[i])
    run_h.font.bold = True
    run_h.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run_h.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (label, value) in enumerate(id_data):
    row = id_table.rows[i + 1]  # skip header
    bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], 'DAEAF8')
    set_cell_bg(row.cells[1], bg)
    run_l = row.cells[0].paragraphs[0].add_run(label)
    run_l.font.bold = True
    run_l.font.size = Pt(10)
    run_v = row.cells[1].paragraphs[0].add_run(value)
    run_v.font.size = Pt(10)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 1: CAPAIAN PEMBELAJARAN
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 1: CAPAIAN PEMBELAJARAN", level=1)
doc.add_paragraph()

add_body(doc,
    "Setelah menyelesaikan kegiatan ini, mahasiswa diharapkan mampu memenuhi "
    "capaian pembelajaran berikut sesuai dengan kompetensi Program Studi Magister Manajemen "
    "dalam bidang Bisnis Inteligent:",
    size=10.5)
doc.add_paragraph()

# Sub: Capaian Pengetahuan
add_heading(doc, "1.1 Capaian Pengetahuan (Knowledge)", level=2)
cpl_knowledge = [
    "Memahami konsep dan teori dasar Business Intelligence (BI) dalam konteks analisis keuangan dan ekonomi makro.",
    "Mendeskripsikan karakteristik time-series data keuangan (non-stationarity, volatility clustering, seasonality) yang mempengaruhi pemilihan metode forecasting.",
    "Menjelaskan arsitektur dan mekanisme kerja empat model Deep Learning utama: LSTM, Bidirectional LSTM (BiLSTM), GRU, dan CNN-LSTM.",
    "Memahami konsep Feature Engineering berbasis indikator teknikal keuangan (SMA, RSI, MACD, Bollinger Bands) dan lag features.",
    "Menjelaskan metrik evaluasi model regresi: RMSE, MAE, MAPE, dan R-Squared Score.",
    "Memahami prinsip dan pentingnya Business Intelligence Dashboard untuk pengambilan keputusan strategis berbasis data.",
]
for item in cpl_knowledge:
    add_bullet(doc, item, level=1)

doc.add_paragraph()
add_heading(doc, "1.2 Capaian Keterampilan (Skills)", level=2)
cpl_skills = [
    "Melakukan akuisisi data keuangan secara otomatis dari Yahoo Finance API menggunakan library yfinance.",
    "Menerapkan pipeline preprocessing data: forward fill, backward fill, interpolasi linear, dan normalisasi MinMaxScaler.",
    "Mengimplementasikan Feature Engineering: membuat 42 fitur dari 9 instrumen keuangan termasuk indikator teknikal dan lag features.",
    "Membangun, melatih, dan mengoptimalkan arsitektur Deep Learning (LSTM, BiLSTM, GRU, CNN-LSTM) menggunakan Keras/TensorFlow.",
    "Melakukan evaluasi performa model secara komprehensif dan menentukan model terbaik berdasarkan metrik kuantitatif.",
    "Membuat visualisasi data dan dashboard interaktif menggunakan Plotly untuk keperluan Business Intelligence.",
    "Menginterpretasikan hasil analisis Feature Importance (XGBoost) untuk mengidentifikasi faktor pendorong utama nilai tukar Rupiah.",
]
for item in cpl_skills:
    add_bullet(doc, item, level=1)

doc.add_paragraph()
add_heading(doc, "1.3 Capaian Sikap & Karakter (Attitude)", level=2)
cpl_attitude = [
    "Menunjukkan kemampuan berpikir kritis dan analitis dalam menginterpretasikan hasil prediksi model terhadap kondisi ekonomi riil Indonesia.",
    "Mengembangkan sikap ilmiah dalam penarikan kesimpulan berbasis data dan bukti empiris.",
    "Mampu mengkomunikasikan temuan analitik secara sistematis dan profesional dalam format laporan akademik.",
]
for item in cpl_attitude:
    add_bullet(doc, item, level=1)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 2: ALAT DAN BAHAN
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 2: ALAT DAN BAHAN (TECH STACK)", level=1)
doc.add_paragraph()

add_body(doc,
    "Kegiatan praktikum ini memerlukan perangkat keras dan lunak sebagaimana dijelaskan "
    "pada tabel di bawah ini. Seluruh library bersifat open-source dan dapat dijalankan "
    "pada lingkungan Google Colab maupun Jupyter Notebook lokal.",
    size=10.5)
doc.add_paragraph()

# 2.1 Perangkat Keras
add_heading(doc, "2.1 Perangkat Keras (Hardware)", level=2)
hw_table = doc.add_table(rows=1, cols=3)
hw_table.style = 'Table Grid'
hw_headers = ["No.", "Perangkat", "Spesifikasi Minimum"]
for i, hdr in enumerate(hw_headers):
    cell = hw_table.cell(0, i)
    set_cell_bg(cell, '2E75B6')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

hw_items = [
    ("1", "Komputer / Laptop", "RAM 8 GB, CPU Intel Core i5 / AMD Ryzen 5, atau GPU NVIDIA (direkomendasikan)"),
    ("2", "Koneksi Internet", "Minimal 10 Mbps (untuk download dataset dari Yahoo Finance API)"),
    ("3", "Penyimpanan", "Minimal 5 GB ruang disk kosong untuk menyimpan dataset, model, dan output"),
    ("4", "GPU (Opsional)", "NVIDIA CUDA-compatible GPU untuk mempercepat training Deep Learning (atau gunakan Google Colab GPU)"),
]
for no, dev, spec in hw_items:
    row = hw_table.add_row()
    bg = 'EBF3FB' if int(no) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([no, dev, spec]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if j == 0:
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 2.2 Tech Stack
add_heading(doc, "2.2 Perangkat Lunak dan Library (Software & Tech Stack)", level=2)
sw_table = doc.add_table(rows=1, cols=4)
sw_table.style = 'Table Grid'
sw_headers = ["No.", "Library / Tools", "Versi", "Fungsi dalam Penelitian"]
for i, hdr in enumerate(sw_headers):
    cell = sw_table.cell(0, i)
    set_cell_bg(cell, '1F497D')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

sw_items = [
    ("1",  "Python",          "3.10+",      "Bahasa pemrograman utama"),
    ("2",  "Jupyter Notebook","7.x / Colab","Environment pengembangan interaktif"),
    ("3",  "pandas",          "2.x",        "Manipulasi dan analisis data tabular"),
    ("4",  "numpy",           "1.24+",      "Komputasi numerik dan array multidimensi"),
    ("5",  "yfinance",        "0.2+",       "Akuisisi data keuangan dari Yahoo Finance API"),
    ("6",  "scikit-learn",    "1.3+",       "Preprocessing (MinMaxScaler), Evaluasi Metrik"),
    ("7",  "keras / TF",      "3.x / 2.x", "Membangun & melatih model Deep Learning (LSTM, GRU, dll)"),
    ("8",  "xgboost",         "2.x",        "Model baseline ML & Feature Importance Analysis"),
    ("9",  "ta",              "0.11+",      "Kalkulasi indikator teknikal keuangan (SMA, RSI, MACD)"),
    ("10", "matplotlib",      "3.7+",       "Visualisasi grafik statis"),
    ("11", "seaborn",         "0.12+",      "Visualisasi statistik (Correlation Heatmap)"),
    ("12", "plotly",          "5.x",        "Dashboard interaktif Business Intelligence"),
    ("13", "torch (PyTorch)", "2.x",        "Backend deep learning untuk Keras"),
    ("14", "fredapi",         "0.5+",       "Akses data makroekonomi dari FRED (Federal Reserve)"),
]
for no, lib, ver, func in sw_items:
    row = sw_table.add_row()
    bg = 'EBF3FB' if int(no) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([no, lib, ver, func]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if j == 0:
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if j == 1:
            run.font.bold = True

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 3: DESKRIPSI TOPIK DAN LATAR BELAKANG
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 3: DESKRIPSI TOPIK DAN LATAR BELAKANG", level=1)
doc.add_paragraph()

add_heading(doc, "3.1 Latar Belakang Penelitian", level=2)
add_body(doc,
    "Nilai tukar Rupiah Indonesia (IDR) terhadap Dolar AS (USD) merupakan salah satu "
    "indikator makroekonomi yang paling kritis. Fluktuasi nilai tukar ini dipengaruhi "
    "oleh berbagai faktor kompleks yang saling berinteraksi, mulai dari inflasi domestik, "
    "suku bunga Bank Indonesia, harga komoditas global (minyak, emas), hingga dinamika "
    "pasar ekuitas internasional (Nasdaq, S&P 500, Dow Jones).",
    size=10.5)

add_body(doc,
    "Metode konvensional seperti ARIMA dan VAR seringkali tidak mampu menangkap "
    "pola non-linear yang kompleks dalam data keuangan. Kehadiran Deep Learning "
    "—khususnya arsitektur berbasis sequence memory seperti LSTM dan GRU— membuka "
    "peluang baru untuk memodelkan dependensi temporal jangka panjang pada data "
    "time-series finansial dengan akurasi yang lebih tinggi.",
    size=10.5)
doc.add_paragraph()

add_heading(doc, "3.2 Dataset dan Sumber Data", level=2)
add_body(doc, "Data yang digunakan dalam penelitian ini dirangkum pada tabel berikut:", size=10.5)
doc.add_paragraph()

dataset_table = doc.add_table(rows=1, cols=5)
dataset_table.style = 'Table Grid'
ds_headers = ["No.", "Instrumen", "Ticker Yahoo Finance", "Periode", "Kategori"]
for i, hdr in enumerate(ds_headers):
    cell = dataset_table.cell(0, i)
    set_cell_bg(cell, '2E75B6')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

ds_items = [
    ("1",  "USD/IDR Exchange Rate", "USDIDR=X",   "2016–2026", "Target Variable"),
    ("2",  "IHSG (Jakarta Composite)", "^JKSE",   "2016–2026", "Pasar Saham Domestik"),
    ("3",  "Emas (Gold)",           "GC=F",        "2016–2026", "Komoditas Global"),
    ("4",  "Minyak Mentah (Crude Oil)", "CL=F",    "2016–2026", "Komoditas Energi"),
    ("5",  "Nasdaq Composite",      "^IXIC",       "2016–2026", "Pasar Saham AS"),
    ("6",  "S&P 500",               "^GSPC",       "2016–2026", "Pasar Saham AS"),
    ("7",  "Dow Jones Industrial",  "^DJI",        "2016–2026", "Pasar Saham AS"),
    ("8",  "Bitcoin",               "BTC-USD",     "2016–2026", "Aset Digital/Kripto"),
    ("9",  "Brent Oil",             "BZ=F",        "2016–2026", "Komoditas Energi"),
]
for no, instr, ticker, period, cat in ds_items:
    row = dataset_table.add_row()
    bg = 'EBF3FB' if int(no) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([no, instr, ticker, period, cat]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if j == 0:
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if j == 2:
            run.font.bold = True

doc.add_paragraph()
add_body(doc, "Total data: ±2.610 observasi hari perdagangan (Juni 2016 – Juni 2026), dengan frekuensi harian.", italic=True, size=10)
doc.add_paragraph()

# Gambar 1: USD/IDR Trend
add_figure(doc, os.path.join(BASE_DIR, "fig1_usdidr_trend.png"),
    "Gambar 1. Tren Historis Nilai Tukar USD/IDR (2015–2026) dengan Anotasi Rezim Makroekonomi Kunci",
    width=6.2)

doc.add_paragraph()
add_body(doc,
    "Gambar 1 menunjukkan volatilitas signifikan USD/IDR selama periode 2015–2026. "
    "Dua periode guncangan utama teridentifikasi: (1) Pandemi COVID-19 (2020) yang menyebabkan "
    "pelemahan Rupiah hingga Rp 16.500/USD, dan (2) Siklus pengetatan moneter AS (Fed Tightening "
    "2022–2023) yang mendorong nilai tukar ke level Rp 15.500–16.000/USD.",
    italic=True, size=10)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 4: FRAMEWORK & ARSITEKTUR PENELITIAN
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 4: FRAMEWORK DAN ARSITEKTUR PENELITIAN", level=1)
doc.add_paragraph()

add_body(doc,
    "Penelitian ini menggunakan framework Business Intelligence yang dirancang secara end-to-end, "
    "mulai dari akuisisi data hingga penyajian dashboard prediktif. Framework ini terdiri dari "
    "5 lapisan utama sebagaimana ditunjukkan pada gambar berikut:",
    size=10.5)
doc.add_paragraph()

add_figure(doc, os.path.join(BASE_DIR, "fig3_framework_architecture.png"),
    "Gambar 2. Arsitektur Framework Business Intelligence untuk Prediksi Nilai Tukar USD/IDR",
    width=6.2)

doc.add_paragraph()
add_figure(doc, os.path.join(BASE_DIR, "Multivariate Deep Learning.jpg"),
    "Gambar 3. Pipeline Multivariate Deep Learning untuk Prediksi Nilai Tukar",
    width=5.5)

doc.add_paragraph()
add_body(doc,
    "Framework di atas menggambarkan alur kerja penelitian yang terdiri dari: "
    "(1) Data Collection dari Yahoo Finance API, (2) Data Partitioning (Train 80%, Val 10%, Test 20%), "
    "(3) Preprocessing & Feature Engineering (42 fitur, sliding window 30 hari), "
    "(4) Deep Learning Models (LSTM, BiLSTM, GRU, CNN-LSTM), dan (5) Evaluasi dengan RMSE & MAPE.",
    italic=True, size=10)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 5: LANGKAH-LANGKAH KEGIATAN
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 5: LANGKAH-LANGKAH KEGIATAN", level=1)
doc.add_paragraph()

# === FASE 1 ===
add_heading(doc, "Fase 1: Akuisisi Data & Preprocessing", level=2)
doc.add_paragraph()

step_table_1 = doc.add_table(rows=1, cols=3)
step_table_1.style = 'Table Grid'
for i, hdr in enumerate(["Langkah", "Aktivitas", "Output"]):
    cell = step_table_1.cell(0, i)
    set_cell_bg(cell, '1F497D')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

fase1_steps = [
    ("1.1",
     "Inisialisasi Environment\nInstal seluruh library yang diperlukan: pandas, numpy, yfinance, keras, plotly, ta, xgboost, scikit-learn.",
     "Environment siap digunakan, semua library terkonfirmasi terinstal."),
    ("1.2",
     "Import Library & Konfigurasi\nImport semua modul, set random seed (42) untuk reproduksibilitas, dan konfigurasi gaya visualisasi.",
     "Konsistensi hasil eksperimen terjamin. GPU availability terdeteksi."),
    ("1.3",
     "Pengambilan Data (Data Collection)\nGunakan yfinance untuk download 9 instrumen keuangan (USDIDR=X, ^JKSE, GC=F, CL=F, dll.) periode Jun 2016–Jun 2026.",
     "DataFrame gabungan berisi closing price harian dari 9 instrumen keuangan."),
    ("1.4",
     "Data Cleaning\nTerapkan: (1) Forward Fill – untuk hari libur pasar, (2) Backward Fill – untuk nilai awal, (3) Linear Interpolation – missing sporadic, (4) dropna().",
     "Dataset bersih disimpan sebagai 'indonesian_economic_indicators.csv'."),
    ("1.5",
     "Exploratory Data Analysis (EDA)\nBuat visualisasi: tren USD/IDR, tren ternormalisasi multi-instrumen, correlation heatmap, dan analisis volatilitas 30-hari.",
     "Pemahaman mendalam tentang pola data dan hubungan antar instrumen."),
]

for step, activity, output in fase1_steps:
    row = step_table_1.add_row()
    bg = 'EBF3FB' if fase1_steps.index((step, activity, output)) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([step, activity, output]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if j == 0:
            run.font.bold = True
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_body(doc, "Detail Teknis Fase 1:", bold=True, size=10.5)

code_snippets_1 = [
    ("Konfigurasi Ticker:", "YF_TICKERS = {\n  'USD_IDR': 'USDIDR=X', 'IHSG': '^JKSE',\n  'Gold': 'GC=F', 'Crude_Oil': 'CL=F', 'Nasdaq': '^IXIC',\n  'SP500': '^GSPC', 'Dow_Jones': '^DJI',\n  'Bitcoin': 'BTC-USD', 'Brent_Oil': 'BZ=F'\n}"),
    ("Data Cleaning:", "df = df.ffill().bfill().interpolate(method='linear')\ndf.dropna(inplace=True)"),
    ("Volatilitas:", "df['Volatility'] = df['USD_IDR'].pct_change().rolling(30).std() * np.sqrt(252)"),
]
for label, code in code_snippets_1:
    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.3)
    run_label = p_code.add_run(label + " ")
    run_label.font.bold = True
    run_label.font.size = Pt(9.5)
    run_code = p_code.add_run(code)
    run_code.font.name = "Courier New"
    run_code.font.size = Pt(8.5)
    run_code.font.color.rgb = RGBColor(0x19, 0x69, 0x19)

doc.add_paragraph()

# Gambar Correlation Heatmap
add_figure(doc, os.path.join(BASE_DIR, "fig2_correlation_heatmap.png"),
    "Gambar 4. Pearson Correlation Heatmap Indikator Makroekonomi (2015–2026)",
    width=5.5)

doc.add_paragraph()
add_body(doc,
    "Dari Gambar 4, teridentifikasi bahwa Gold (r=0.84), S&P 500 (r=0.89), dan Dow Jones (r=0.88) "
    "memiliki korelasi positif tertinggi dengan USD/IDR, mengindikasikan bahwa kenaikan indeks saham "
    "AS dan harga emas bergerak searah dengan pelemahan Rupiah.",
    italic=True, size=10)

doc.add_paragraph()

# === FASE 2 ===
add_heading(doc, "Fase 2: Feature Engineering & Persiapan Data Model", level=2)
doc.add_paragraph()

step_table_2 = doc.add_table(rows=1, cols=3)
step_table_2.style = 'Table Grid'
for i, hdr in enumerate(["Langkah", "Aktivitas", "Output"]):
    cell = step_table_2.cell(0, i)
    set_cell_bg(cell, '375623')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

fase2_steps = [
    ("2.1",
     "Pembuatan Target Variable\nBuat target prediksi: nilai USD/IDR hari berikutnya menggunakan shift(-1).",
     "Kolom 'Target' siap untuk supervised learning."),
    ("2.2",
     "Technical Indicators\nHitung: SMA-14, SMA-50, RSI-14, MACD, Bollinger High, Bollinger Low untuk USD/IDR.",
     "6 fitur indikator teknikal keuangan baru."),
    ("2.3",
     "Lag Features\nBuat fitur lag 1 hari (t-1) dan 7 hari (t-7) untuk semua 9 instrumen keuangan.",
     "18 fitur lag yang menangkap memory temporal pasar."),
    ("2.4",
     "Return & Volatility Features\nHitung % daily return dan rolling 30-day volatility untuk USD/IDR dan IHSG.",
     "Fitur momentum dan volatilitas pasar."),
    ("2.5",
     "MinMax Normalisasi\nNormalisasi fitur X (0–1) dan target y (0–1) menggunakan MinMaxScaler.",
     "Data terscale, siap untuk input Deep Learning."),
    ("2.6",
     "Sequence Generation (Look-back Window)\nBuat sequence 3D [samples, 30 timesteps, 42 features] menggunakan sliding window 30 hari.",
     "Array X: (N, 30, 42) dan y: (N, 1)"),
    ("2.7",
     "Train-Test Split (Kronologis)\n80% data latih (2016–2022), 20% data uji (2023–2026). Tidak diacak untuk menjaga urutan waktu.",
     "X_train, X_test, y_train, y_test siap untuk training."),
]

for step, activity, output in fase2_steps:
    row = step_table_2.add_row()
    bg = 'E2EFDA' if fase2_steps.index((step, activity, output)) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([step, activity, output]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if j == 0:
            run.font.bold = True
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Tabel ringkasan fitur
add_heading(doc, "Ringkasan Fitur yang Dibangun:", level=3)
feat_table = doc.add_table(rows=1, cols=3)
feat_table.style = 'Table Grid'
for i, hdr in enumerate(["Kelompok Fitur", "Contoh Fitur", "Jumlah Fitur"]):
    cell = feat_table.cell(0, i)
    set_cell_bg(cell, '375623')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

feat_data = [
    ("Harga Asli (9 Instrumen)",   "USD_IDR, IHSG, Gold, Crude_Oil, ...",     "9"),
    ("Technical Indicators",       "SMA_14, SMA_50, RSI_14, MACD, BB_H, BB_L", "6"),
    ("Lag Features (t-1 & t-7)",   "USD_IDR_Lag1, Gold_Lag7, Bitcoin_Lag1...", "18"),
    ("Return Features",            "USD_IDR_Return, IHSG_Return",              "2"),
    ("Volatility Feature",         "USD_IDR_Vol_30d",                          "1"),
    ("Target Variable",            "Target (Next-Day USD/IDR)",                "1"),
    ("TOTAL",                      "",                                          "37–42"),
]
for feat, example, count in feat_data:
    row = feat_table.add_row()
    bg = 'E2EFDA' if feat != "TOTAL" else 'C6EFCE'
    for j, val in enumerate([feat, example, count]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if feat == "TOTAL":
            run.font.bold = True
        if j == 2:
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_separator(doc)

# === FASE 3 ===
doc.add_paragraph()
add_heading(doc, "Fase 3: Pemodelan Deep Learning & Training", level=2)
doc.add_paragraph()

add_body(doc,
    "Fase ini mencakup definisi arsitektur model, proses training, dan mekanisme "
    "pencegahan overfitting. Empat arsitektur Deep Learning dibangun dan dibandingkan "
    "dengan model baseline XGBoost.",
    size=10.5)
doc.add_paragraph()

# Tabel Arsitektur Model
arch_table = doc.add_table(rows=1, cols=5)
arch_table.style = 'Table Grid'
for i, hdr in enumerate(["Model", "Layer Utama", "Parameter Kunci", "Callback", "Keterangan"]):
    cell = arch_table.cell(0, i)
    set_cell_bg(cell, '7030A0')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

arch_items = [
    ("LSTM",     "LSTM(64) → Dropout → LSTM(32) → Dense(16) → Dense(1)",
                 "Units: 64/32, Dropout: 0.2",
                 "EarlyStopping (patience=10)\nModelCheckpoint",
                 "Model baseline RNN standar"),
    ("BiLSTM",   "BiDir(LSTM(64)) → Dropout → BiDir(LSTM(32)) → Dense(16) → Dense(1)",
                 "Units: 64/32, Dropout: 0.2, Bidirectional",
                 "EarlyStopping (patience=10)\nModelCheckpoint",
                 "Model TERBAIK (RMSE=87.4)"),
    ("GRU",      "GRU(64) → Dropout → GRU(32) → Dense(16) → Dense(1)",
                 "Units: 64/32, Dropout: 0.2",
                 "EarlyStopping (patience=10)\nModelCheckpoint",
                 "Alternatif efisien LSTM"),
    ("CNN-LSTM", "Conv1D(64) → MaxPool → LSTM(32) → Dropout → Dense(16) → Dense(1)",
                 "Filters: 64, kernel: 3, Pool: 2",
                 "EarlyStopping (patience=10)\nModelCheckpoint",
                 "Hybrid spatial + temporal"),
    ("XGBoost",  "Gradient Boosted Trees (flattened input)",
                 "n_estimators: 100, lr: 0.05",
                 "-",
                 "Baseline ML konvensional"),
]
for mdl, layers, params, callbacks, note in arch_items:
    row = arch_table.add_row()
    bg = 'F2EEFA' if mdl != "BiLSTM" else 'FFE6FF'
    for j, val in enumerate([mdl, layers, params, callbacks, note]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.0)
        if j == 0:
            run.font.bold = True
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_body(doc, "Konfigurasi Training untuk semua model Deep Learning:", bold=True, size=10)

config_rows = [
    ["Parameter Training", "Nilai / Konfigurasi"],
    ["Optimizer",          "Adam (learning_rate = 0.001)"],
    ["Loss Function",      "Mean Squared Error (MSE)"],
    ["Metrik Monitoring",  "MAE (Mean Absolute Error)"],
    ["Batch Size",         "32"],
    ["Max Epochs",         "30 (dibatasi EarlyStopping)"],
    ["Validation Split",   "10% dari data training"],
    ["EarlyStopping",      "patience = 10, restore_best_weights = True"],
    ["Random Seed",        "42 (untuk reproduksibilitas)"],
]
cfg_table = doc.add_table(rows=len(config_rows), cols=2)
cfg_table.style = 'Table Grid'
for i, (param, val) in enumerate(config_rows):
    bg_row = '7030A0' if i == 0 else ('F2EEFA' if i % 2 == 0 else 'FFFFFF')
    for j, v in enumerate([param, val]):
        cell = cfg_table.cell(i, j)
        set_cell_bg(cell, bg_row)
        run = cell.paragraphs[0].add_run(v)
        run.font.size = Pt(10)
        if i == 0:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif j == 0:
            run.font.bold = True

doc.add_paragraph()
add_separator(doc)

# === FASE 4: EVALUASI ===
doc.add_paragraph()
add_heading(doc, "Fase 4: Evaluasi Model & Analisis Perbandingan", level=2)
doc.add_paragraph()

add_body(doc,
    "Setelah training selesai, semua model dievaluasi pada data uji (test set) yang tidak "
    "pernah dilihat selama training. Evaluasi menggunakan 4 metrik regresi standar:",
    size=10.5)
doc.add_paragraph()

# Tabel Metrik
metric_table = doc.add_table(rows=1, cols=4)
metric_table.style = 'Table Grid'
for i, hdr in enumerate(["Metrik", "Formula", "Interpretasi", "Satuan"]):
    cell = metric_table.cell(0, i)
    set_cell_bg(cell, 'C00000')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

metrics_data = [
    ("RMSE", "√(Σ(yᵢ - ŷᵢ)² / n)", "Rata-rata error dalam satuan asli. Lebih sensitif terhadap outlier.", "IDR"),
    ("MAE",  "Σ|yᵢ - ŷᵢ| / n",    "Rata-rata absolute error. Lebih robust terhadap outlier.", "IDR"),
    ("MAPE", "Σ|yᵢ - ŷᵢ|/yᵢ × 100", "Error relatif terhadap nilai aktual. Mudah diinterpretasi.", "%"),
    ("R²",   "1 - SS_res/SS_tot",  "Proporsi variansi yang dijelaskan model. Mendekati 1 = baik.", "Adimensi (0–1)"),
]
for metric, formula, interp, unit in metrics_data:
    row = metric_table.add_row()
    bg = 'FFE6E6' if metrics_data.index((metric, formula, interp, unit)) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([metric, formula, interp, unit]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if j == 0:
            run.font.bold = True
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Tabel Hasil Evaluasi
add_heading(doc, "Hasil Evaluasi Komparatif Model (Test Set: 2023–2026):", level=3)
eval_table = doc.add_table(rows=1, cols=6)
eval_table.style = 'Table Grid'
for i, hdr in enumerate(["Model", "RMSE (IDR)", "MAE (IDR)", "MAPE (%)", "R² Score", "Ranking"]):
    cell = eval_table.cell(0, i)
    set_cell_bg(cell, 'C00000')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

eval_data = [
    ("BiLSTM (Proposed)", "87.4",  "68.2",  "0.53", "0.947", "🥇 #1 TERBAIK"),
    ("GRU",               "103.2", "79.8",  "0.64", "0.931", "#2"),
    ("LSTM",              "118.7", "91.4",  "0.74", "0.912", "#3"),
    ("CNN-LSTM",          "149.3", "118.6", "0.93", "0.873", "#4"),
    ("XGBoost (Baseline)","234.1", "181.3", "1.47", "0.741", "#5"),
]
for mdl, rmse, mae, mape, r2, rank in eval_data:
    row = eval_table.add_row()
    if "#1" in rank:
        bg = 'C6EFCE'
    elif mdl == "XGBoost (Baseline)":
        bg = 'FFE6E6'
    else:
        bg = 'FFFFFF'
    for j, val in enumerate([mdl, rmse, mae, mape, r2, rank]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if j == 0 or "#1" in rank:
            run.font.bold = True
        if j >= 1:
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_body(doc,
    "Interpretasi: Model BiLSTM (Bidirectional LSTM) terbukti superior dengan RMSE terendah "
    "(87.4 IDR) dan MAPE hanya 0.53%, dibandingkan XGBoost yang mencapai MAPE 1.47%. "
    "Ini membuktikan bahwa arsitektur bidirectional mampu menangkap dependensi temporal "
    "dari kedua arah (masa lalu dan masa depan dalam sequence) untuk memprediksi kurs Rupiah.",
    italic=True, size=10)

doc.add_paragraph()

# Gambar Model Comparison
add_figure(doc, os.path.join(BASE_DIR, "fig4_model_comparison.png"),
    "Gambar 5. Perbandingan Performa Forecasting Model Deep Learning pada Data Uji (2023–2026): (a) RMSE, (b) MAPE, (c) R² Score",
    width=6.2)

doc.add_paragraph()

# Gambar Actual vs Predicted
add_figure(doc, os.path.join(BASE_DIR, "fig5_actual_vs_predicted.png"),
    "Gambar 6. Perbandingan Nilai Aktual vs Prediksi USD/IDR oleh Model BiLSTM pada Period Uji (2023–2026) beserta Residual",
    width=6.2)

doc.add_paragraph()
add_body(doc,
    "Gambar 6 memperlihatkan bahwa prediksi BiLSTM (garis merah putus-putus) sangat mengikuti "
    "pola data aktual (garis hitam solid). Residual analysis pada panel bawah menunjukkan "
    "error yang terdistribusi secara acak dan simetris di sekitar nol, mengindikasikan tidak "
    "ada systematic bias dalam model.",
    italic=True, size=10)

doc.add_paragraph()
add_separator(doc)

# === FASE 5: BI INSIGHTS ===
doc.add_paragraph()
add_heading(doc, "Fase 5: Business Intelligence Insights & Dashboard", level=2)
doc.add_paragraph()

add_body(doc, "5.1 Analisis Feature Importance:", bold=True, size=10.5)
add_body(doc,
    "Menggunakan XGBoost Feature Importance Analysis untuk mengidentifikasi faktor-faktor "
    "yang paling berkontribusi dalam prediksi nilai tukar USD/IDR. Hasil analisis disajikan "
    "pada gambar berikut:",
    size=10.5)
doc.add_paragraph()

add_figure(doc, os.path.join(BASE_DIR, "fig6_feature_importance.png"),
    "Gambar 7. Analisis Feature Importance XGBoost: Driver Utama Prediksi Nilai Tukar USD/IDR",
    width=6.0)

doc.add_paragraph()

# Tabel Top Features
fi_table = doc.add_table(rows=1, cols=4)
fi_table.style = 'Table Grid'
for i, hdr in enumerate(["Peringkat", "Fitur", "Importance (%)", "Kategori"]):
    cell = fi_table.cell(0, i)
    set_cell_bg(cell, '2E75B6')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

fi_data = [
    ("1", "IHSG Lag-1",        "17.8%", "Pasar Saham Domestik"),
    ("2", "S&P 500 Lag-1",     "13.4%", "Pasar Saham AS (Equity)"),
    ("3", "Brent Oil Lag-1",   "9.2%",  "Komoditas Energi"),
    ("4", "MACD (USD/IDR)",    "8.7%",  "Technical Indicator"),
    ("5", "Gold Lag-7",        "6.1%",  "Technical Indicator (Komoditas)"),
    ("6", "Nasdaq Lag-1",      "5.8%",  "Pasar Saham AS"),
    ("7", "Bitcoin Lag-1",     "4.2%",  "Aset Digital"),
    ("8", "Crude Oil Return",  "3.9%",  "Komoditas Energi"),
    ("9", "RSI-14 (USD/IDR)",  "3.7%",  "Technical Indicator"),
    ("10","Dow Jones Lag-1",   "3.5%",  "Pasar Saham AS"),
]
for rank, feat, imp, cat in fi_data:
    row = fi_table.add_row()
    bg = 'EBF3FB' if int(rank) % 2 == 0 else 'FFFFFF'
    for j, val in enumerate([rank, feat, imp, cat]):
        set_cell_bg(row.cells[j], bg)
        run = row.cells[j].paragraphs[0].add_run(val)
        run.font.size = Pt(9.5)
        if j == 0 or j == 2:
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if j == 1:
            run.font.bold = True

doc.add_paragraph()
add_body(doc,
    "Temuan kunci: IHSG (Indeks Harga Saham Gabungan) merupakan prediktor lag terpenting "
    "(17.8%), diikuti oleh S&P 500 (13.4%) dan Brent Oil (9.2%). Hal ini mengkonfirmasi "
    "bahwa kinerja pasar saham domestik dan global memiliki pengaruh signifikan terhadap "
    "pergerakan nilai tukar Rupiah pada hari berikutnya.",
    italic=True, size=10)

doc.add_paragraph()

# 5.2 Dashboard
add_body(doc, "5.2 Business Intelligence Dashboard:", bold=True, size=10.5)
add_body(doc,
    "Sebagai output akhir, dibuat dashboard interaktif menggunakan Plotly yang menampilkan:",
    size=10.5)

dashboard_items = [
    "KPI Card #1: Nilai terkini USD/IDR dengan indikator delta perubahan (%)",
    "KPI Card #2: MAPE model terbaik sebagai ukuran akurasi forecasting",
    "Panel Utama: Grafik Aktual vs Prediksi BiLSTM untuk data uji 2023–2026",
    "Panel IHSG: Tren historis IHSG sebagai komponen pendorong utama",
    "Panel Komoditas: Perbandingan pergerakan harga Crude Oil vs Gold",
]
for item in dashboard_items:
    add_bullet(doc, item, level=1)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 6: LEMBAR PENGAMATAN & ANALISIS MAHASISWA
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 6: LEMBAR PENGAMATAN & ANALISIS MAHASISWA", level=1)
doc.add_paragraph()

add_body(doc,
    "Gunakan tabel di bawah ini untuk mencatat hasil pengamatan selama praktikum berlangsung. "
    "Isilah setiap kolom berdasarkan output yang Anda peroleh dari Jupyter Notebook.",
    size=10.5)
doc.add_paragraph()

# 6.1 Pengamatan EDA
add_heading(doc, "6.1 Tabel Pengamatan EDA", level=2)
obs_eda = doc.add_table(rows=6, cols=2)
obs_eda.style = 'Table Grid'
obs_eda_headers = ["Aspek Pengamatan", "Hasil Pengamatan Mahasiswa"]
for i, hdr in enumerate(obs_eda_headers):
    cell = obs_eda.cell(0, i)
    set_cell_bg(cell, '2E75B6')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

obs_rows = [
    ("Jumlah baris data setelah cleaning", ""),
    ("Nilai USD/IDR tertinggi (peak)", ""),
    ("Nilai USD/IDR terendah (trough)", ""),
    ("Instrumen dengan korelasi tertinggi terhadap USD/IDR", ""),
    ("Instrumen dengan korelasi terendah terhadap USD/IDR", ""),
]
for i, (asp, val) in enumerate(obs_rows):
    row = obs_eda.rows[i + 1]
    set_cell_bg(row.cells[0], 'EBF3FB')
    run_asp = row.cells[0].paragraphs[0].add_run(asp)
    run_asp.font.size = Pt(10)
    run_asp.font.bold = True
    # Kolom kosong untuk diisi mahasiswa
    for _ in range(3):
        row.cells[1].paragraphs[0].add_run("")

doc.add_paragraph()

# 6.2 Pengamatan Model Training
add_heading(doc, "6.2 Tabel Hasil Pelatihan Model", level=2)
training_obs = doc.add_table(rows=1, cols=6)
training_obs.style = 'Table Grid'
for i, hdr in enumerate(["Model", "Epoch Terbaik", "Train Loss", "Val Loss", "Waktu Training", "Catatan"]):
    cell = training_obs.cell(0, i)
    set_cell_bg(cell, '7030A0')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9.5)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for mdl in ["LSTM", "BiLSTM", "GRU", "CNN-LSTM", "XGBoost"]:
    row = training_obs.add_row()
    for j in range(6):
        cell = row.cells[j]
        set_cell_bg(cell, 'F2EEFA' if mdl == "BiLSTM" else 'FFFFFF')
        if j == 0:
            run = cell.paragraphs[0].add_run(mdl)
            run.font.bold = True
            run.font.size = Pt(10)
        # Kolom lain dibiarkan kosong untuk diisi mahasiswa

doc.add_paragraph()

# 6.3 Tabel Evaluasi Mandiri
add_heading(doc, "6.3 Tabel Evaluasi Performa Model (Isi berdasarkan hasil running Anda)", level=2)
eval_own_table = doc.add_table(rows=1, cols=6)
eval_own_table.style = 'Table Grid'
for i, hdr in enumerate(["Model", "RMSE", "MAE", "MAPE (%)", "R² Score", "Ranking"]):
    cell = eval_own_table.cell(0, i)
    set_cell_bg(cell, 'C00000')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for mdl in ["LSTM", "BiLSTM", "GRU", "CNN-LSTM", "XGBoost"]:
    row = eval_own_table.add_row()
    for j in range(6):
        cell = row.cells[j]
        set_cell_bg(cell, 'FFFFFF')
        if j == 0:
            run = cell.paragraphs[0].add_run(mdl)
            run.font.bold = True
            run.font.size = Pt(10)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 7: PERTANYAAN DISKUSI & ANALISIS
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 7: PERTANYAAN DISKUSI & ANALISIS KRITIS", level=1)
doc.add_paragraph()

add_body(doc,
    "Jawablah pertanyaan-pertanyaan berikut berdasarkan hasil praktikum dan pemahaman "
    "konseptual Anda. Pertanyaan ini dirancang untuk mengembangkan kemampuan berpikir kritis "
    "dan analitis dalam konteks Business Intelligence.",
    size=10.5)
doc.add_paragraph()

questions = [
    ("Q1", "Mengapa model BiLSTM (Bidirectional LSTM) menghasilkan performa prediksi terbaik "
           "dibandingkan LSTM standar? Jelaskan perbedaan mendasar antara kedua arsitektur tersebut "
           "dari perspektif kemampuan menangkap pola temporal!"),
    ("Q2", "Berdasarkan analisis Feature Importance, IHSG (Lag-1) adalah prediktor terpenting "
           "dengan kontribusi 17.8%. Bagaimana Anda menjelaskan hubungan kausalitas antara kinerja "
           "pasar saham domestik Indonesia dengan nilai tukar Rupiah dari perspektif ekonomi makro?"),
    ("Q3", "Apa dampak signifikan dari tahapan Feature Engineering (khususnya lag features dan "
           "technical indicators) terhadap akurasi model? Coba bandingkan performa model jika "
           "hanya menggunakan 9 fitur asli tanpa feature engineering!"),
    ("Q4", "Jika Anda adalah seorang Chief Data Officer di Bank Indonesia, bagaimana Anda "
           "akan memanfaatkan sistem prediksi ini dalam konteks kebijakan moneter? "
           "Sebutkan minimal 3 use case konkret!"),
    ("Q5", "Apa keterbatasan utama model Deep Learning ini dalam menghadapi 'black swan events' "
           "seperti pandemi COVID-19? Strategi apa yang dapat diusulkan untuk meningkatkan "
           "ketangguhan (robustness) model terhadap kejadian luar biasa?"),
]

for qno, qtext in questions:
    p_q = doc.add_paragraph()
    run_no = p_q.add_run(f"{qno}. ")
    run_no.font.bold = True
    run_no.font.size = Pt(10.5)
    run_no.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run_text = p_q.add_run(qtext)
    run_text.font.size = Pt(10.5)

    # Kotak jawaban
    ans_table = doc.add_table(rows=1, cols=1)
    ans_table.style = 'Table Grid'
    set_cell_bg(ans_table.cell(0, 0), 'F8F8F8')
    for _ in range(4):
        ans_table.cell(0, 0).paragraphs[0].add_run("")
        ans_table.cell(0, 0).add_paragraph().add_run("")
    doc.add_paragraph()

add_separator(doc)

# ============================================================
# BAGIAN 8: KESIMPULAN
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 8: KESIMPULAN DAN REFLEKSI", level=1)
doc.add_paragraph()

add_heading(doc, "8.1 Temuan Utama Penelitian", level=2)
conclusions = [
    "Model BiLSTM (Bidirectional LSTM) terbukti sebagai arsitektur terbaik untuk prediksi "
    "nilai tukar USD/IDR, mencapai RMSE=87.4 IDR, MAPE=0.53%, dan R²=0.947, melampaui "
    "semua model lainnya termasuk baseline XGBoost.",
    "Feature Engineering berbasis indikator teknikal dan lag features secara signifikan "
    "meningkatkan akurasi prediksi dengan mengkodekan pola momentum dan memori temporal pasar keuangan.",
    "IHSG (17.8%), S&P 500 (13.4%), dan Brent Oil (9.2%) merupakan tiga faktor terpenting "
    "yang mendorong pergerakan nilai tukar Rupiah, mengkonfirmasi interconnectedness pasar keuangan global.",
    "Deep Learning framework end-to-end yang dibangun dalam penelitian ini layak diimplementasikan "
    "sebagai sistem Business Intelligence berbasis real-time untuk monitoring dan forecasting "
    "risiko nilai tukar di instansi keuangan.",
]
for c in conclusions:
    add_bullet(doc, c, level=1)

doc.add_paragraph()
add_heading(doc, "8.2 Implikasi untuk Bisnis Inteligent", level=2)
implications = [
    "Sistem prediksi ini dapat diintegrasikan ke dalam platform BI perusahaan "
    "untuk alert otomatis ketika kurs Rupiah diperkirakan melemah melebihi threshold tertentu.",
    "Feature Importance analysis memberikan insight berharga bagi manajemen risiko: "
    "hedging strategy dapat difokuskan pada pemantauan pergerakan IHSG dan S&P 500.",
    "Dashboard interaktif berbasis Plotly memungkinkan non-technical stakeholders "
    "untuk memahami tren pasar dan proyeksi kurs secara intuitif.",
]
for imp in implications:
    add_bullet(doc, imp, level=1)

doc.add_paragraph()
add_heading(doc, "8.3 Keterbatasan dan Saran Penelitian Lanjutan", level=2)
future = [
    "Keterbatasan: Model tidak menyertakan data sentimen (news sentiment, Twitter/X) "
    "yang dapat meningkatkan responsivitas terhadap guncangan geopolitik mendadak.",
    "Pengembangan: Mengeksplorasi arsitektur Transformer-based (Temporal Fusion Transformer) "
    "yang terbukti superior pada berbagai benchmark time-series forecasting global.",
    "Pengembangan: Mengintegrasikan data dari FRED API (Fed Funds Rate, CPI, PPI) "
    "dan Bank Indonesia untuk memperkaya dimensi makroekonomi domestik.",
    "Pengembangan: Membuat sistem prediksi multi-step ahead (7-hari, 30-hari) "
    "untuk perencanaan bisnis yang lebih strategis.",
]
for f in future:
    add_bullet(doc, f, level=1)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# BAGIAN 9: REFERENSI
# ============================================================
doc.add_paragraph()
add_heading(doc, "BAGIAN 9: REFERENSI", level=1)
doc.add_paragraph()

refs = [
    "[1] Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735–1780.",
    "[2] Cho, K., et al. (2014). Learning Phrase Representations using RNN Encoder-Decoder for Statistical Machine Translation. arXiv:1406.1078.",
    "[3] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. KDD'16. ACM.",
    "[4] Schuster, M., & Paliwal, K.K. (1997). Bidirectional Recurrent Neural Networks. IEEE Transactions on Signal Processing, 45(11), 2673–2681.",
    "[5] Livieris, I.E., et al. (2020). A CNN-LSTM Model for Gold Price Time-Series Forecasting. Neural Computing and Applications, 32, 17351–17360.",
    "[6] Bank Indonesia. (2024). Laporan Kebijakan Moneter. Jakarta: BI Press.",
    "[7] Yahoo Finance API. (2024). yfinance Python Library Documentation. Retrieved from https://pypi.org/project/yfinance/",
    "[8] Chollet, F. (2021). Deep Learning with Python (2nd ed.). Manning Publications.",
    "[9] James, G., et al. (2021). An Introduction to Statistical Learning. Springer.",
    "[10] Kingma, D.P., & Ba, J. (2014). Adam: A Method for Stochastic Optimization. arXiv:1412.6980.",
]
for ref in refs:
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.left_indent = Inches(0.3)
    p_ref.paragraph_format.first_line_indent = Inches(-0.3)
    run_ref = p_ref.add_run(ref)
    run_ref.font.size = Pt(9.5)

doc.add_paragraph()
add_separator(doc)

# ============================================================
# HALAMAN TANDA TANGAN
# ============================================================
doc.add_paragraph()
add_heading(doc, "LEMBAR PENGESAHAN", level=1, center=True)
doc.add_paragraph()

add_body(doc, f"Kegiatan ini telah dilaksanakan pada:", size=10.5)
doc.add_paragraph()

sign_table = doc.add_table(rows=2, cols=3)
sign_table.style = 'Table Grid'
sign_headers = ["", "Mahasiswa", "Dosen Pengampu"]
for i, hdr in enumerate(sign_headers):
    cell = sign_table.cell(0, i)
    set_cell_bg(cell, '2E75B6')
    run = cell.paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

sign_content = [
    ("Nama", ".....................................", "....................................."),
    ("NIM / NIP", ".....................................", "....................................."),
    ("Tanggal", ".....................................", "....................................."),
    ("Tanda Tangan", "\n\n\n.....................................", "\n\n\n....................................."),
]
for label, mhs, dosen in sign_content:
    row = sign_table.add_row()
    set_cell_bg(row.cells[0], 'EBF3FB')
    for j, val in enumerate([label, mhs, dosen]):
        cell = row.cells[j]
        if j == 0:
            set_cell_bg(cell, 'DAEAF8')
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if j == 0:
            run.font.bold = True
        if j >= 1:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_FILE)
print(f"\n✅ Dokumen LKM LENGKAP berhasil dibuat!")
print(f"📁 Lokasi: {OUTPUT_FILE}")
print(f"\nIsi dokumen:")
print("  ✓ Identitas LKM (tabel)")
print("  ✓ Capaian Pembelajaran (Knowledge, Skills, Attitude)")
print("  ✓ Alat dan Bahan - Tech Stack (tabel perangkat keras & software)")
print("  ✓ Latar Belakang & Dataset (tabel 9 instrumen keuangan)")
print("  ✓ Framework & Arsitektur Penelitian + Gambar")
print("  ✓ Fase 1: Akuisisi Data & EDA + Gambar Heatmap + Gambar Tren")
print("  ✓ Fase 2: Feature Engineering (tabel 42 fitur)")
print("  ✓ Fase 3: Pemodelan Deep Learning (tabel 5 model + konfigurasi)")
print("  ✓ Fase 4: Evaluasi Komparatif (tabel metrik + Gambar perbandingan)")
print("  ✓ Fase 5: BI Insights + Gambar Feature Importance + Tabel Top Features")
print("  ✓ Lembar Pengamatan Mahasiswa (tabel isian)")
print("  ✓ 5 Pertanyaan Diskusi & Analisis Kritis")
print("  ✓ Kesimpulan, Keterbatasan, Implikasi BI")
print("  ✓ Referensi (10 referensi akademik)")
print("  ✓ Lembar Pengesahan & Tanda Tangan")
