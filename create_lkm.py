import os

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    print("Menginstal python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Judul
title = doc.add_paragraph()
run = title.add_run("LEMBAR KEGIATAN MAHASISWA (LKM)")
run.font.bold = True
run.font.size = Pt(14)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Informasi Umum
doc.add_paragraph("Mata Kuliah: Bisnis Inteligent").runs[0].font.bold = True
doc.add_paragraph("Topik: Analisis dan Prediksi Nilai Tukar Rupiah Terhadap Dolar AS Menggunakan Deep Learning").runs[0].font.bold = True

doc.add_paragraph()

# 1. Capaian Pembelajaran
p1 = doc.add_paragraph()
r1 = p1.add_run("1. Capaian Pembelajaran")
r1.font.bold = True

doc.add_paragraph("• Mahasiswa mampu memahami konsep dasar pemodelan time-series menggunakan arsitektur Deep Learning (LSTM, BiLSTM, GRU, CNN-LSTM).")
doc.add_paragraph("• Mahasiswa mampu melakukan proses akuisisi data keuangan dan ekonomi (Indikator Makroekonomi, Nilai Tukar) dari sumber API publik (Yahoo Finance).")
doc.add_paragraph("• Mahasiswa mampu melakukan tahapan preprocessing data, termasuk penanganan missing values, normalisasi, dan pembuatan fitur (Feature Engineering) berbasis indikator teknikal keuangan.")
doc.add_paragraph("• Mahasiswa mampu mengimplementasikan dan melatih model Deep Learning untuk peramalan (forecasting) nilai tukar uang.")
doc.add_paragraph("• Mahasiswa mampu mengevaluasi dan membandingkan performa berbagai model menggunakan metrik standar (RMSE, MAE, MAPE, R²).")
doc.add_paragraph("• Mahasiswa mampu menyajikan wawasan analitik ekonomi secara interaktif (Business Intelligence Dashboard) menggunakan Plotly.")

doc.add_paragraph()

# 2. Alat dan Bahan (Tech Stack)
p2 = doc.add_paragraph()
r2 = p2.add_run("2. Alat dan Bahan (Tech Stack)")
r2.font.bold = True

doc.add_paragraph("• Bahasa Pemrograman: Python 3")
doc.add_paragraph("• Environment: Jupyter Notebook / Google Colab")
doc.add_paragraph("• Data Manipulation & Analysis: pandas, numpy")
doc.add_paragraph("• Data Visualization & Dashboarding: matplotlib, seaborn, plotly")
doc.add_paragraph("• Financial Data APIs: yfinance, fredapi")
doc.add_paragraph("• Machine Learning & Deep Learning: scikit-learn, keras, tensorflow, xgboost")
doc.add_paragraph("• Technical Analysis Indicators: ta")

doc.add_paragraph()

# 3. Langkah-Langkah Kegiatan
p3 = doc.add_paragraph()
r3 = p3.add_run("3. Langkah-Langkah Kegiatan")
r3.font.bold = True

# Fase 1
p3_1 = doc.add_paragraph()
r3_1 = p3_1.add_run("Fase 1: Akuisisi Data & Preprocessing")
r3_1.font.bold = True

doc.add_paragraph("1. Inisialisasi Environment: Menginstal semua dependensi dan library yang dibutuhkan seperti yfinance, keras, pandas, plotly, dll.")
doc.add_paragraph("2. Import Library: Memuat library untuk manipulasi data (pandas, numpy), visualisasi (matplotlib, plotly), pengambilan data (yfinance), preprocessing (scikit-learn), dan pemodelan (keras, xgboost).")
doc.add_paragraph("3. Pengambilan Data Historis (Data Collection):")
doc.add_paragraph("   - Menggunakan `yfinance` untuk mengunduh data historis (Juni 2016 - Juni 2026) untuk instrumen keuangan: USD/IDR, IHSG, Gold, Crude Oil, Nasdaq, S&P 500, Dow Jones, Bitcoin, dan Brent Oil.")
doc.add_paragraph("   - Menggabungkan data seluruh instrumen dalam satu DataFrame terpusat berdasarkan indeks tanggal waktu (datetime).")
doc.add_paragraph("4. Data Cleaning (Pembersihan Data):")
doc.add_paragraph("   - Melakukan forward fill (ffill) untuk mengatasi hari-hari libur (weekends/holidays) ketika pasar tutup.")
doc.add_paragraph("   - Melakukan backward fill (bfill) dan interpolasi linear untuk menangani missing values.")
doc.add_paragraph("   - Menghapus nilai NaN yang masih tersisa (dropna) untuk mendapatkan dataset bersih ('indonesian_economic_indicators.csv').")
doc.add_paragraph("5. Exploratory Data Analysis (EDA):")
doc.add_paragraph("   - Membuat visualisasi tren nilai tukar USD/IDR.")
doc.add_paragraph("   - Menganalisis perbandingan tren yang telah dinormalisasi antara USD/IDR, IHSG, Crude Oil, dan Gold.")
doc.add_paragraph("   - Mengkaji korelasi antar berbagai indikator ekonomi menggunakan Correlation Heatmap.")
doc.add_paragraph("   - Menganalisis tingkat volatilitas historis (rolling 30-day standard deviation) dari USD/IDR.")

doc.add_paragraph()

# Fase 2
p3_2 = doc.add_paragraph()
r3_2 = p3_2.add_run("Fase 2: Feature Engineering & Persiapan Data Model")
r3_2.font.bold = True

doc.add_paragraph("1. Pembuatan Fitur (Feature Engineering):")
doc.add_paragraph("   - Target Variable: Mempersiapkan target prediksi yaitu nilai USD/IDR pada periode selanjutnya (shift -1).")
doc.add_paragraph("   - Technical Indicators: Menghitung Simple Moving Average (SMA), Relative Strength Index (RSI), MACD, dan Bollinger Bands untuk USD/IDR.")
doc.add_paragraph("   - Lag Features: Membuat fitur lag (hari sebelumnya) untuk menangkap pola berulang dari seluruh instrumen keuangan.")
doc.add_paragraph("   - Return & Volatility: Menghitung persentase perubahan harga dan volatilitas per 30 hari untuk analisis momentum pasar.")
doc.add_paragraph("2. Scaling Data: Melakukan normalisasi fitur input (X) dan target (y) menggunakan MinMaxScaler sehingga nilainya berada di antara skala (0, 1), cocok untuk Deep Learning.")
doc.add_paragraph("3. Sequence Generation: Membuat format time-series berurutan dengan look-back window 30 hari (memprediksi hari ke-31 berdasarkan 30 hari ke belakang).")
doc.add_paragraph("4. Train-Test Split: Membagi dataset menjadi data latih (80%) dan data uji (20%) secara kronologis (tidak diacak).")

doc.add_paragraph()

# Fase 3
p3_3 = doc.add_paragraph()
r3_3 = p3_3.add_run("Fase 3: Pemodelan Deep Learning & Evaluasi")
r3_3.font.bold = True

doc.add_paragraph("1. Arsitektur Model: Mengembangkan dan melatih empat jenis arsitektur Deep Learning:")
doc.add_paragraph("   - LSTM (Long Short-Term Memory)")
doc.add_paragraph("   - Bidirectional LSTM (BiLSTM)")
doc.add_paragraph("   - GRU (Gated Recurrent Unit)")
doc.add_paragraph("   - CNN-LSTM (Kombinasi Convolutional 1D dengan LSTM)")
doc.add_paragraph("2. Mekanisme Training: Menggunakan optimizer Adam, Loss function MSE, dilengkapi callbacks EarlyStopping dan ModelCheckpoint untuk menghindari overfitting.")
doc.add_paragraph("3. Baseline Model: Melatih model Machine Learning konvensional (XGBoost Regressor) sebagai pembanding performa.")
doc.add_paragraph("4. Evaluasi & Metrik: Menghitung RMSE, MAE, MAPE, dan R-Squared dari hasil prediksi tiap model terhadap data uji untuk menemukan model terbaik.")

doc.add_paragraph()

# Fase 4
p3_4 = doc.add_paragraph()
r3_4 = p3_4.add_run("Fase 4: Visualisasi BI Dashboard & Penarikan Insight")
r3_4.font.bold = True

doc.add_paragraph("1. Business Intelligence Report: Mengekstraksi temuan terkait faktor dominan (korelasi positif/negatif terbesar), pentingnya fitur (Feature Importance dari XGBoost), dan superioritas arsitektur terpilih.")
doc.add_paragraph("2. Pembuatan Dashboard: Menyusun dan merender dasbor interaktif menggunakan Plotly yang menampilkan KPI (Nilai Terbaru, MAPE Model Terbaik), Grafik Prediksi vs Aktual, Tren IHSG, serta harga Komoditas (Oil vs Gold).")
doc.add_paragraph("3. Penyusunan Laporan Otomatis: Menyimpulkan hasil prediksi dan keterkaitannya dengan indikator ekonomi dalam format ringkasan naratif.")

output_file = "LKM_Bisnis_Inteligent.docx"
doc.save(output_file)
print(f"Dokumen LKM berhasil dibuat di {os.path.abspath(output_file)}")
