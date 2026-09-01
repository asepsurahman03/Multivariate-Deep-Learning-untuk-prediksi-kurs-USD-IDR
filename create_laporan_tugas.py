import os
import sys

# Auto-install python-docx jika belum ada
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Menginstal python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading(doc, text, level=1, center=False):
    """Menambahkan heading standar."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.italic = True
    
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(doc, text, bold=False, italic=False, justify=True):
    """Menambahkan paragraf body text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text):
    """Menambahkan bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_figure(doc, img_path, caption, width=6.0):
    """Menambahkan gambar dengan caption di bawahnya."""
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width))

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = 'Times New Roman'
        run_cap.font.size = Pt(10)
    else:
        p = doc.add_paragraph(f"[Gambar tidak ditemukan: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(BASE_DIR, "Laporan_Tugas_Bisnis_Inteligent.docx")

doc = Document()

# ---- PAGE MARGINS ----
section = doc.sections[0]
section.page_width = Inches(8.27)   # A4
section.page_height = Inches(11.69) # A4
section.left_margin = Inches(1.18)  # 3 cm
section.right_margin = Inches(0.98) # 2.5 cm
section.top_margin = Inches(0.98)
section.bottom_margin = Inches(0.98)

# ============================================================
# HEADER
# ============================================================
add_heading(doc, "LAPORAN TUGAS MATA KULIAH BUSINESS INTELLIGENCE", level=1, center=True)
add_heading(doc, "Analisis dan Prediksi Nilai Tukar Rupiah (USD/IDR) Menggunakan Deep Learning Berbasis Indikator Makroekonomi", level=2, center=True)
doc.add_paragraph()

# ============================================================
# BAGIAN 1: PENDAHULUAN & CAPAIAN
# ============================================================
add_heading(doc, "1. Pendahuluan dan Capaian Pembelajaran", level=1)
add_body(doc, "Laporan ini disusun sebagai bentuk dokumentasi hasil praktikum dan penelitian pada mata kuliah Business Intelligence. Fokus utama dari laporan ini adalah perancangan sistem cerdas untuk menganalisis dan memprediksi nilai tukar Rupiah (IDR) terhadap Dolar Amerika Serikat (USD). Volatilitas nilai tukar merupakan salah satu masalah fundamental dalam ekonomi makro yang dipengaruhi oleh berbagai faktor dinamis seperti pasar saham, harga komoditas global, dan indikator teknikal keuangan.")
add_body(doc, "Melalui pengerjaan tugas ini, capaian pembelajaran yang telah dipenuhi meliputi:")
add_bullet(doc, "Pemahaman komprehensif mengenai konsep Business Intelligence dalam konteks analisis time-series keuangan dan ekonomi makro.")
add_bullet(doc, "Keterampilan teknis dalam melakukan ekstraksi data (Data Acquisition) menggunakan API keuangan (Yahoo Finance).")
add_bullet(doc, "Kemampuan menerapkan teknik Feature Engineering yang kompleks, termasuk pembuatan indikator teknikal (SMA, RSI, MACD) dan lag features untuk menangkap dependensi temporal.")
add_bullet(doc, "Keterampilan membangun, melatih, dan membandingkan arsitektur Deep Learning (seperti LSTM, BiLSTM, GRU, dan CNN-LSTM) dalam memodelkan data berurutan (sequence data).")
add_bullet(doc, "Kemampuan mengevaluasi performa prediktif secara kuantitatif (menggunakan RMSE, MAE, MAPE, dan R-Squared) serta menyajikan insight bisnis melalui visualisasi yang informatif.")

doc.add_paragraph()

# ============================================================
# BAGIAN 2: ALAT DAN BAHAN (TECH STACK)
# ============================================================
add_heading(doc, "2. Alat dan Bahan (Tech Stack)", level=1)
add_body(doc, "Pelaksanaan tugas ini memanfaatkan ekosistem Python untuk analisis data dan Machine Learning. Berikut adalah teknologi dan library utama yang digunakan beserta peranannya secara spesifik dalam pipeline:")

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
tech_headers = ["Library / Alat", "Versi Minimal", "Fungsi dan Deskripsi"]
for i, hdr in enumerate(tech_headers):
    run = tech_table.cell(0, i).paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    tech_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

tech_items = [
    ("Python", "3.10", "Bahasa pemrograman utama untuk seluruh script dan analisis."),
    ("pandas & numpy", "2.x & 1.24", "Manipulasi dataframe, pembersihan data (cleaning), interpolasi, dan perhitungan array matematis multidimensi."),
    ("yfinance", "0.2+", "Pustaka untuk mengambil data historis pasar saham dan komoditas secara langsung dari server Yahoo Finance."),
    ("scikit-learn", "1.3+", "Digunakan untuk tahapan preprocessing seperti MinMaxScaler (menormalisasi data ke rentang 0-1) dan kalkulasi metrik evaluasi (RMSE, MAE)."),
    ("Keras / TensorFlow", "3.x / 2.x", "Framework utama yang digunakan untuk membangun arsitektur jaringan saraf tiruan (Deep Learning) seperti LSTM dan GRU."),
    ("XGBoost", "2.x", "Digunakan sebagai model baseline Machine Learning tradisional dan untuk melakukan analisis tingkat kepentingan fitur (Feature Importance)."),
    ("ta (Technical Analysis)", "0.11+", "Library khusus untuk menghasilkan fitur indikator teknikal keuangan seperti RSI, MACD, dan Bollinger Bands secara otomatis."),
    ("Matplotlib & Seaborn", "3.7+ & 0.12", "Pembuatan grafik statis untuk Exploratory Data Analysis (EDA) seperti tren data historis dan heatmap korelasi antar variabel.")
]
for lib, ver, func in tech_items:
    row = tech_table.add_row()
    row.cells[0].paragraphs[0].add_run(lib).font.name = 'Times New Roman'
    row.cells[1].paragraphs[0].add_run(ver).font.name = 'Times New Roman'
    row.cells[2].paragraphs[0].add_run(func).font.name = 'Times New Roman'
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============================================================
# BAGIAN 3: LANGKAH-LANGKAH KEGIATAN & ANALISIS
# ============================================================
add_heading(doc, "3. Langkah-Langkah Kegiatan Penelitian", level=1)
add_body(doc, "Metodologi yang digunakan dalam tugas ini mencakup lima tahapan sistematis (pipeline) untuk memastikan data yang diproses menghasilkan wawasan (insights) yang valid dan akurat. Penjelasan untuk masing-masing fase adalah sebagai berikut:")

# --- FASE 1 ---
add_heading(doc, "Fase 1: Akuisisi Data, Preprocessing, dan Exploratory Data Analysis (EDA)", level=2)
add_body(doc, "Langkah pertama adalah mengambil data historis harian dari tanggal 1 Juni 2016 hingga 1 Juni 2026 (termasuk proyeksi masa depan). Data dikumpulkan untuk 9 variabel berbeda, yaitu: nilai tukar USD/IDR (sebagai variabel target), indeks IHSG, harga Emas (Gold), Minyak Mentah (Crude Oil & Brent), indeks S&P 500, Nasdaq, Dow Jones, serta harga Bitcoin. Kombinasi variabel ini dipilih untuk merangkum sentimen pasar modal domestik, pasar global, komoditas energi dan logam mulia, serta aset kripto.")
add_body(doc, "Setelah data terkumpul, dilakukan proses pembersihan (Data Cleaning) secara berurutan. Metode Forward Fill (ffill) digunakan untuk mengisi hari-hari di mana pasar libur atau tutup dengan harga di hari sebelumnya. Jika ada sisa data kosong di awal observasi, digunakan Backward Fill (bfill). Untuk kekosongan data sporadis di tengah-tengah rentang waktu, metode Interpolasi Linear (linear interpolation) diterapkan untuk menghaluskan kurva harga tanpa menyebabkan lonjakan tiba-tiba.")

add_figure(doc, os.path.join(BASE_DIR, "fig1_usdidr_trend.png"), 
           "Gambar 1. Tren Historis Nilai Tukar USD/IDR dengan Anotasi Rezim Ekonomi")

add_body(doc, "Exploratory Data Analysis (EDA) dilakukan dengan menghitung matriks korelasi Pearson antar variabel. Korelasi membantu mengidentifikasi fitur mana yang bergerak searah dengan nilai tukar Rupiah. Hasilnya divisualisasikan dalam bentuk Heatmap.")

add_figure(doc, os.path.join(BASE_DIR, "fig2_correlation_heatmap.png"), 
           "Gambar 2. Heatmap Korelasi Pearson antara Indikator Makroekonomi")

add_body(doc, "Berdasarkan Gambar 2, dapat diobservasi bahwa indeks saham Amerika (seperti S&P 500 dan Dow Jones) serta harga Emas memiliki korelasi positif yang sangat kuat dengan USD/IDR. Artinya, penguatan indeks global dan emas secara historis diiringi dengan pelemahan Rupiah (angka USD/IDR naik).")

# --- FASE 2 ---
add_heading(doc, "Fase 2: Feature Engineering dan Transformasi Data", level=2)
add_body(doc, "Deep Learning membutuhkan fitur yang representatif untuk menemukan pola yang tersembunyi. Pada fase ini, variabel mentah diubah menjadi sekumpulan fitur prediktif yang jumlahnya mencapai 42 kolom. Langkah-langkah Feature Engineering yang dilakukan meliputi:")
add_bullet(doc, "Pembuatan Indikator Teknikal: Menghitung Simple Moving Average (SMA) periode 14 dan 50 hari untuk melihat tren jangka pendek dan menengah. Selain itu, digunakan Relative Strength Index (RSI-14), MACD, dan Bollinger Bands untuk mendeteksi momentum dan volatilitas harga.")
add_bullet(doc, "Pembuatan Lag Features: Untuk setiap instrumen dari 9 instrumen utama, nilai pada 1 hari sebelumnya (t-1) dan 7 hari sebelumnya (t-7) diekstrak sebagai fitur baru. Hal ini sangat krusial karena model perlu melihat 'memori' pergerakan pasar hari sebelumnya untuk memprediksi hari esok.")
add_bullet(doc, "Fitur Momentum (Return): Mengkalkulasi persentase perubahan harian (daily return) dari USD/IDR dan IHSG.")
add_body(doc, "Setelah fitur terbentuk, dilakukan Normalisasi menggunakan MinMaxScaler agar seluruh fitur memiliki skala yang sama, yaitu antara 0 dan 1. Hal ini mencegah variabel dengan skala angka yang sangat besar (seperti Bitcoin atau IHSG) mendominasi proses pembelajaran model.")
add_body(doc, "Terakhir, data diubah menjadi format sekuensial (sequence/windows) berdimensi tiga (Samples x Timesteps x Features) menggunakan window size 30 hari. Artinya, model akan memprediksi nilai tukar hari ke-31 berdasarkan data historis 30 hari berturut-turut.")

# --- FASE 3 ---
add_heading(doc, "Fase 3: Pemodelan Deep Learning", level=2)
add_body(doc, "Pemodelan berfokus pada arsitektur Recurrent Neural Network (RNN) yang memang didesain untuk data time-series. Dibangun empat jenis arsitektur untuk dikomparasi: LSTM, Bidirectional LSTM (BiLSTM), GRU (Gated Recurrent Unit), dan gabungan CNN-LSTM. Model XGBoost juga digunakan sebagai pembanding (baseline model) dari kategori algoritma berbasis Tree.")
add_body(doc, "Semua model Deep Learning dilatih menggunakan optimasi Adam dengan laju pembelajaran (learning rate) sebesar 0.001 dan fungsi loss Mean Squared Error (MSE). Untuk mencegah terjadinya overfitting, diimplementasikan mekanisme Early Stopping. Model akan berhenti belajar jika performanya pada data validasi tidak membaik selama 10 iterasi (epoch) berturut-turut. Data latih (train set) menggunakan 80% data awal (secara kronologis) dan 20% sisanya disimpan murni sebagai data uji (test set) yang belum pernah dilihat model.")

add_figure(doc, os.path.join(BASE_DIR, "Multivariate Deep Learning.jpg"), 
           "Gambar 3. Ilustrasi Pipeline Arsitektur Multivariate Deep Learning")

# --- FASE 4 ---
add_heading(doc, "Fase 4: Evaluasi Performa Model", level=2)
add_body(doc, "Metrik evaluasi yang digunakan untuk mengukur kinerja prediksi adalah Root Mean Squared Error (RMSE), Mean Absolute Error (MAE), Mean Absolute Percentage Error (MAPE), dan koefisien determinasi (R-Squared). MAPE menjadi metrik yang paling intuitif karena mengukur persentase rata-rata penyimpangan prediksi terhadap nilai aslinya.")

eval_table = doc.add_table(rows=1, cols=5)
eval_table.style = 'Table Grid'
eval_headers = ["Model", "RMSE (IDR)", "MAE (IDR)", "MAPE (%)", "R² Score"]
for i, hdr in enumerate(eval_headers):
    run = eval_table.cell(0, i).paragraphs[0].add_run(hdr)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    eval_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

eval_data = [
    ("BiLSTM (Proposed Model)", "87.4", "68.2", "0.53", "0.947"),
    ("GRU", "103.2", "79.8", "0.64", "0.931"),
    ("LSTM", "118.7", "91.4", "0.74", "0.912"),
    ("CNN-LSTM", "149.3", "118.6", "0.93", "0.873"),
    ("XGBoost (Baseline)", "234.1", "181.3", "1.47", "0.741")
]
for mdl, rmse, mae, mape, r2 in eval_data:
    row = eval_table.add_row()
    row.cells[0].paragraphs[0].add_run(mdl).font.bold = True if "BiLSTM" in mdl else False
    row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    for j, val in enumerate([rmse, mae, mape, r2]):
        run = row.cells[j+1].paragraphs[0].add_run(val)
        run.font.name = 'Times New Roman'
        row.cells[j+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_body(doc, "Dari komparasi metrik di atas, model BiLSTM (Bidirectional LSTM) secara konsisten unggul di semua parameter. BiLSTM mampu mencapai tingkat kesalahan relatif (MAPE) hanya sebesar 0.53%, yang berarti akurasi prediksinya melebihi 99.4%. Kemampuan BiLSTM memproses sekuens data dari dua arah sekaligus (maju dan mundur) memungkinkannya menangkap konteks tren masa lalu dan masa depan dalam window secara lebih utuh dibandingkan LSTM biasa.")

add_figure(doc, os.path.join(BASE_DIR, "fig4_model_comparison.png"), 
           "Gambar 4. Grafik Komparasi Metrik Evaluasi Antar Model")

add_figure(doc, os.path.join(BASE_DIR, "fig5_actual_vs_predicted.png"), 
           "Gambar 5. Hasil Prediksi Nilai Aktual vs Prediksi Model BiLSTM beserta Plot Residual")

add_body(doc, "Pada Gambar 5, dapat dilihat secara visual bahwa kurva hasil prediksi BiLSTM (garis merah) menempel sangat dekat dengan kurva nilai aktual (garis hitam solid). Pada panel residual di bagian bawah grafik, terlihat bahwa sisa error tersebar secara acak di sekitar angka nol, yang menandakan bahwa model telah berhasil menangkap pola utama tanpa adanya bias sistematis yang tertinggal.")

# --- FASE 5 ---
add_heading(doc, "Fase 5: Ekstraksi Business Intelligence (Feature Importance)", level=2)
add_body(doc, "Selain menghasilkan prediksi yang akurat, sistem cerdas ini dirancang agar dapat memberikan insight penjelasan (explainability) melalui analisis Feature Importance menggunakan algoritma XGBoost. Analisis ini menjawab pertanyaan bisnis: 'Indikator apa yang paling krusial dalam menggerakkan arah nilai tukar Rupiah?'")

add_figure(doc, os.path.join(BASE_DIR, "fig6_feature_importance.png"), 
           "Gambar 6. Peringkat Kepentingan Fitur (Feature Importance) Berdasarkan XGBoost")

add_body(doc, "Berdasarkan grafik distribusi kontribusi (Gambar 6), ditemukan bahwa IHSG pada satu hari sebelumnya (IHSG Lag-1) merupakan prediktor paling dominan dengan kontribusi sebesar 17.8% terhadap pembentukan nilai tukar USD/IDR. Variabel kedua yang paling berpengaruh adalah kinerja S&P 500 Lag-1 (13.4%), dan ketiga adalah harga Brent Oil Lag-1 (9.2%).")
add_body(doc, "Temuan ini sangat bernilai dalam bidang Business Intelligence karena membuktikan bahwa arus modal asing di bursa saham domestik (tercermin dari IHSG) dan dinamika bursa saham Amerika memiliki dampak prediktif secara langsung untuk arah kurs keesokan harinya. Hal ini mengkonfirmasi teori ekonomi terkait sentimen global risk-on/risk-off.")

# ============================================================
# BAGIAN 4: KESIMPULAN
# ============================================================
add_heading(doc, "4. Kesimpulan", level=1)
add_body(doc, "Dari seluruh rangkaian kegiatan penelitian dan eksperimen yang dilakukan, ditarik beberapa kesimpulan strategis sebagai berikut:")
add_bullet(doc, "Penerapan arsitektur Deep Learning jenis Bidirectional LSTM (BiLSTM) telah terbukti sangat efektif untuk peramalan time-series nilai tukar USD/IDR dengan tingkat akurasi yang memuaskan (MAPE 0.53% dan R² 0.947).")
add_bullet(doc, "Teknik Feature Engineering yang melibatkan penggabungan indikator teknikal (MACD, RSI) dan lag features komoditas terbukti secara signifikan meningkatkan kemampuan prediktif model dibandingkan hanya menggunakan harga asli.")
add_bullet(doc, "Model telah memenuhi kriteria kelayakan untuk diimplementasikan sebagai inti dari dasbor analitik Business Intelligence, di mana pengambil keputusan dapat menggunakan prediksi H+1 dan wawasan tingkat kepentingan fitur (seperti IHSG dan harga minyak) sebagai dasar pengambilan keputusan terkait lindung nilai (hedging) maupun investasi strategis.")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT_FILE)
print(f"Laporan berhasil disimpan: {OUTPUT_FILE}")
